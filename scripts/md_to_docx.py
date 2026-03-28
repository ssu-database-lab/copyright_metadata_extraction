#!/usr/bin/env python3
"""
Convert API_SPECIFICATION.md to DOCX for sharing.
Open the resulting DOCX in Hancom Office (한글) and Save As → HWP to get HWP format.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, start, end (values like 1 for single line)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'bottom', 'start', 'end'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)


def add_table_from_markdown(doc, lines, start_idx):
    """Parse markdown table and add to document. Returns end index."""
    # Find table boundaries
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [c.strip() for c in lines[i].split('|')[1:-1]]
        if not all(c.replace('-', '').strip() == '' for c in row):  # Skip separator row
            rows.append(row)
        i += 1
    
    if not rows:
        return start_idx
    
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < len(table.rows[ri].cells):
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
    return i


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert Markdown to DOCX with formatting."""
    doc = Document()
    
    # Set default font for Korean
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    
    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()
        
        # Code block
        if line_stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                code_lang = line_stripped[3:].strip()
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Headers
        if line_stripped.startswith('## '):
            p = doc.add_heading(line_stripped[3:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        if line_stripped.startswith('### '):
            p = doc.add_heading(line_stripped[4:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        if line_stripped.startswith('#### '):
            p = doc.add_heading(line_stripped[5:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Horizontal rule
        if line_stripped in ('---', '***', '___'):
            doc.add_paragraph('_' * 60)
            i += 1
            continue
        
        # Table
        if line_stripped.startswith('|') and '|' in line_stripped:
            i = add_table_from_markdown(doc, lines, i)
            continue
        
        # Bold list items
        if line_stripped.startswith('- **') and '**:' in line_stripped:
            match = re.match(r'^- \*\*(.+?)\*\*:?\s*(.*)$', line_stripped)
            if match:
                p = doc.add_paragraph()
                run1 = p.add_run(match.group(1) + ': ')
                run1.bold = True
                run2 = p.add_run(match.group(2))
                p.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        
        # List items
        if line_stripped.startswith('- '):
            p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        
        # Empty line
        if not line_stripped:
            i += 1
            continue
        
        # Regular paragraph (handle **bold**)
        p = doc.add_paragraph()
        remaining = line_stripped
        while remaining:
            if '**' in remaining:
                before, mid, after = remaining.split('**', 2)
                p.add_run(before)
                run = p.add_run(mid)
                run.bold = True
                remaining = after
            else:
                p.add_run(remaining)
                break
        p.paragraph_format.space_after = Pt(4)
        i += 1
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    print("To get HWP: Open in Hancom Office (한글) → File → Save As → HWP format")


if __name__ == '__main__':
    base = Path(__file__).parent.parent
    md_path = base / 'API_SPECIFICATION.md'
    docx_path = base / 'API_SPECIFICATION.docx'
    md_to_docx(md_path, docx_path)
